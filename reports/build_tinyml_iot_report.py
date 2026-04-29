from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCREENSHOTS = ROOT / "screenshots"
REPORT_DIR = ROOT / "reports"
DOCX_PATH = REPORT_DIR / "tinyml_iot_demo_report.docx"
ARCH_PATH = SCREENSHOTS / "19_architecture_diagram.png"


ACCENT = RGBColor(28, 103, 88)
MUTED = RGBColor(86, 95, 112)
DARK = RGBColor(18, 30, 47)
LIGHT_FILL = "EAF4F1"
TABLE_FILL = "DCEBE7"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "",
        "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_rounded_box(draw: ImageDraw.ImageDraw, xy, fill, outline, text, text_fill, title=False):
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    f = font(26 if title else 21, bold=title)
    lines = wrap_text(text, f, x2 - x1 - 34)
    y = y1 + 22
    for line in lines:
        draw.text((x1 + 18, y), line, font=f, fill=text_fill)
        y += 31 if title else 27


def wrap_text(text: str, fnt: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    dummy = Image.new("RGB", (1, 1))
    draw = ImageDraw.Draw(dummy)
    for word in words:
        candidate = f"{current} {word}".strip()
        width = draw.textbbox((0, 0), candidate, font=fnt)[2]
        if width <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def arrow(draw: ImageDraw.ImageDraw, start, end, label: str = ""):
    draw.line([start, end], fill=(30, 69, 87), width=4)
    sx, sy = start
    ex, ey = end
    if ex >= sx:
        head = [(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)]
    else:
        head = [(ex, ey), (ex + 14, ey - 8), (ex + 14, ey + 8)]
    draw.polygon(head, fill=(30, 69, 87))
    if label:
        f = font(16)
        mx = (sx + ex) // 2
        my = (sy + ey) // 2 - 22
        bbox = draw.textbbox((mx, my), label, font=f)
        draw.rounded_rectangle((bbox[0] - 8, bbox[1] - 4, bbox[2] + 8, bbox[3] + 4), radius=8, fill=(255, 255, 255))
        draw.text((mx, my), label, font=f, fill=(32, 54, 68), anchor="la")


def generate_architecture_diagram() -> None:
    width, height = 1800, 980
    img = Image.new("RGB", (width, height), (246, 250, 249))
    draw = ImageDraw.Draw(img)
    title_font = font(38, bold=True)
    subtitle_font = font(20)
    draw.text((70, 48), "TinyML IoT demo architecture and data flow", font=title_font, fill=(18, 30, 47))
    draw.text(
        (72, 102),
        "Voice, colour, and movement inference run on the board; session storage and replay run in the web/cloud layer.",
        font=subtitle_font,
        fill=(83, 96, 112),
    )

    draw_rounded_box(
        draw,
        (70, 190, 410, 390),
        (232, 245, 241),
        (28, 103, 88),
        "Arduino Nano 33 BLE Sense\nPDM microphone\nAPDS9960 colour sensor\nLSM9DS1 IMU",
        (18, 30, 47),
        True,
    )
    draw_rounded_box(
        draw,
        (510, 170, 880, 420),
        (255, 255, 255),
        (80, 129, 150),
        "Firmware state machine\nWAITING_FOR_VOICE\nWAITING_FOR_COLOUR\nTRACKING_MOVEMENT\nCombined Edge Impulse inference layer",
        (18, 30, 47),
        True,
    )
    draw_rounded_box(
        draw,
        (980, 200, 1285, 390),
        (255, 255, 255),
        (80, 129, 150),
        "Python serial bridge\nReads JSON lines\nForwards HTTP events\nPolls stop control",
        (18, 30, 47),
        True,
    )
    draw_rounded_box(
        draw,
        (1380, 160, 1725, 430),
        (232, 245, 241),
        (28, 103, 88),
        "Next.js web API\nPOST /api/movement\nGET /api/latest\nGET /api/sessions\nPOST /complete",
        (18, 30, 47),
        True,
    )
    draw_rounded_box(
        draw,
        (1380, 560, 1725, 760),
        (255, 255, 255),
        (80, 129, 150),
        "PostgreSQL in Docker\nsessions\nevents\nmovement_samples\nrecording_state",
        (18, 30, 47),
        True,
    )
    draw_rounded_box(
        draw,
        (980, 570, 1285, 760),
        (255, 255, 255),
        (80, 129, 150),
        "Browser dashboard\nLive monitor\nRecorded sessions\nStop current session\nMovement replay chart",
        (18, 30, 47),
        True,
    )

    arrow(draw, (410, 290), (510, 290), "sensor readings")
    arrow(draw, (880, 290), (980, 290), "USB serial JSON @ 115200")
    arrow(draw, (1285, 290), (1380, 290), "Bearer BRIDGE_API_TOKEN")
    arrow(draw, (1555, 430), (1555, 560), "SQL persistence")
    arrow(draw, (1380, 660), (1285, 660), "public reads + admin stop")

    draw.rounded_rectangle((925, 120, 1335, 470), radius=26, outline=(198, 92, 61), width=3)
    draw.text((948, 130), "Trust boundary: laptop bridge -> HTTP API", font=font(17, bold=True), fill=(120, 61, 42))
    draw.rounded_rectangle((1350, 130, 1755, 795), radius=26, outline=(198, 92, 61), width=3)
    draw.text((1372, 805), "Docker/server boundary", font=font(17, bold=True), fill=(120, 61, 42))

    legend_y = 875
    draw.rounded_rectangle((70, legend_y, 1725, 935), radius=16, fill=(231, 237, 241), outline=(196, 207, 214))
    draw.text(
        (92, legend_y + 17),
        "Security controls: bridge ingest and bridge control use BRIDGE_API_TOKEN; dashboard stop uses ADMIN_API_TOKEN; Postgres is bound to localhost in Docker Compose.",
        font=font(20),
        fill=(18, 30, 47),
    )
    img.save(ARCH_PATH)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(1.75)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.4)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 28, DARK),
        ("Heading 1", 17, ACCENT),
        ("Heading 2", 13, DARK),
        ("Heading 3", 11, ACCENT),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1"} else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.paragraph_format.space_before = Pt(8 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("TinyML IoT Demo Report")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_fill(cell, LIGHT_FILL)
    set_cell_margin(cell, 140, 160, 140, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.size = Pt(10)


def add_table(doc: Document, headers: Iterable[str], rows: Iterable[Iterable[str]]) -> None:
    headers = list(headers)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_fill(cell, TABLE_FILL)
        set_cell_margin(cell, 100, 100, 100, 100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9.2)
        run.font.color.rgb = DARK
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_margin(cell, 90, 100, 90, 100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            run.font.size = Pt(8.8)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 5.9) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(16)
    r = title.add_run("TinyML IoT Demo\nTechnical and Business Report")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = DARK

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(20)
    rs = sub.add_run("Voice recognition personal contribution and group integration report")
    rs.font.size = Pt(13)
    rs.font.color.rgb = MUTED

    add_callout(
        doc,
        "Project scope",
        "Arduino Nano 33 BLE Sense firmware runs three TinyML models for voice, colour authentication, and movement recognition. A Python bridge forwards serial JSON to a Next.js dashboard with Postgres-backed session recording, Docker deployment, and STRIDE-informed API security.",
    )

    meta = [
        ("Student", "Sergey Sotskiy"),
        ("Module", "Internet of Things / Cloud Computing coursework"),
        ("Report split", "Personal section: voice recognition TinyML. Group section: three-model integration, bridge, dashboard, Docker, and security."),
        ("Date", "28 April 2026"),
    ]
    add_table(doc, ["Field", "Value"], meta)
    doc.add_page_break()


def executive_summary(doc: Document) -> None:
    heading(doc, "Executive summary")
    p(
        doc,
        "This report describes a local-first TinyML Internet of Things demo that combines edge inference, cloud-style persistence, and a web dashboard into one demonstrable workflow. The system uses an Arduino Nano 33 BLE Sense class board to detect the spoken keyword start, authenticate a green colour target, and then classify board movement. The board does not connect to the network directly. Instead, it prints structured JSON over USB serial at 115200 baud and a Python bridge forwards each valid event to a Next.js API. The API stores demo sessions in PostgreSQL when Docker is used, while the dashboard shows the latest event, recorded sessions, authentication progress, and a movement replay chart.",
    )
    p(
        doc,
        "The business value is a low-cost, explainable proof of concept for edge sensing workflows. Voice and colour gates make the system behave like an authenticated interaction rather than a raw sensor logger, while movement recognition produces useful session evidence after the user has passed both gates. Running inference on-device reduces dependency on cloud audio processing and keeps the web layer focused on persistence, visibility, and control. Docker Compose makes the dashboard and database repeatable on a server, supporting the cloud-computing requirement without forcing the microcontroller to manage Wi-Fi or database credentials.",
    )
    p(
        doc,
        "My personal contribution was the voice recognition TinyML component. I trained and validated a keyword model in Edge Impulse, assessed its validation and on-device performance, and integrated it into the firmware state machine so a confident start detection advances the system from WAITING_FOR_VOICE to WAITING_FOR_COLOUR. The group contribution merged three Edge Impulse models into one firmware build, built the serial-to-HTTP bridge, implemented session storage and replay, added Docker deployment, and applied a STRIDE-based security layer with bearer-token protected mutation APIs.",
    )


def intro(doc: Document) -> None:
    heading(doc, "1. Introduction and business value")
    p(
        doc,
        "Many IoT demonstrations stop at collecting raw sensor values. This project takes a more product-like approach: the device must first hear a voice command, then authenticate a colour, and only then record movement. The result is a staged interaction that can be explained to non-technical stakeholders and tested live. It demonstrates how embedded machine learning can add local intelligence to a small sensor platform while a web application provides the visibility and evidence trail expected from a connected system.",
    )
    p(
        doc,
        "The Arduino Nano 33 BLE Sense is suitable for this because it combines a microphone, an APDS9960 colour/proximity sensor, and an LSM9DS1 inertial measurement unit with a microcontroller capable of TinyML workloads (Arduino, n.d.-a). Edge Impulse supports training and deploying embedded models as Arduino libraries, making it appropriate for a coursework prototype where the model pipeline and firmware deployment both need to be visible (Edge Impulse, n.d.-a). The system also reflects a practical deployment constraint: the board performs local sensing and inference, but a laptop bridge handles networking and the cloud-facing dashboard.",
    )
    p(
        doc,
        "The business value is not only technical novelty. A staged edge workflow can be adapted to access control, equipment-use logging, gesture-controlled interfaces, or training demonstrations where the system must prove that a user completed a required sequence. The dashboard records sessions, stores movement samples, and can stop a run from the browser. This creates auditable evidence for a demo, while Docker and PostgreSQL make the server side repeatable and persistent rather than a temporary local page.",
    )


def personal_section(doc: Document) -> None:
    heading(doc, "2. Personal contribution: voice recognition TinyML")
    p(
        doc,
        "My personal contribution was the voice recognition stage. The purpose of this model was deliberately narrow: identify the keyword start strongly enough to trigger the next phase, while treating background speech or silence as unknown. This small scope made the model suitable for a constrained embedded device because it avoided open-ended speech recognition and instead used keyword spotting. In the complete system, this model is the first authentication gate, so its output controls whether the demo remains in WAITING_FOR_VOICE or advances to colour authentication.",
    )
    add_figure(
        doc,
        SCREENSHOTS / "personal/voice_dataset.png",
        "Fig. P1. Edge Impulse voice dataset showing the start and unknown classes used to train the keyword model.",
        4.7,
    )
    p(
        doc,
        "The dataset was organised in Edge Impulse around the target command and a negative class. Figure P1 shows the captured voice data used for this model. The positive class represents the word start, while the unknown class gives the classifier examples of audio that should not trigger the state change. This distinction is important because a demo environment can contain speech, movement noise, keyboard sounds, and room noise. A binary keyword model is only useful if it learns both the intended command and the non-command background distribution.",
    )
    add_figure(
        doc,
        SCREENSHOTS / "personal/data_explorer_voice_model.png",
        "Fig. P2. Data explorer view used to inspect recorded voice samples before training.",
        5.8,
    )
    p(
        doc,
        "Before training, the data explorer was used to inspect whether the recorded samples looked plausible and whether the classes were separated enough for a simple audio classifier. Edge Impulse audio workflows normally transform raw audio into signal-processing features before classification, rather than feeding a full waveform directly to an embedded neural network (Edge Impulse, n.d.-b). This approach is appropriate for microcontrollers because it reduces the amount of data passed into the model and makes inference more predictable on constrained hardware. The model therefore fits the TinyML pattern: collect representative data, extract compact features, train a classifier, and deploy the inference code to the target.",
    )
    add_figure(
        doc,
        SCREENSHOTS / "personal/last_training_performance(validation_set).png",
        "Fig. P3. Final validation performance after training the voice model.",
        5.8,
    )
    add_figure(
        doc,
        SCREENSHOTS / "personal/confusion_matrix(validation_set).png",
        "Fig. P4. Confusion matrix showing how validation samples were classified.",
        5.8,
    )
    add_figure(
        doc,
        SCREENSHOTS / "personal/metrics(validation_set)_voice_model.png",
        "Fig. P5. Validation metrics used to judge voice model quality.",
        5.8,
    )
    p(
        doc,
        "The validation results in Figures P3, P4, and P5 were used to decide whether the model was reliable enough for the live state machine. The confusion matrix is particularly useful because the cost of a false positive is high: the system would skip the voice phase and immediately ask for colour authentication. For a demo, occasional false negatives are less damaging because the user can say start again, but false positives make the interaction look uncontrolled. This is why the firmware also includes state-machine protections, including an arming delay after boot and a confidence threshold before the start event is accepted.",
    )
    add_figure(
        doc,
        SCREENSHOTS / "personal/on_device_performace_of_the_voice_model.png",
        "Fig. P6. On-device performance estimate for memory and latency suitability.",
        5.8,
    )
    p(
        doc,
        "On-device performance was checked before integration. Figure P6 shows the deployment estimate used to judge whether the voice model could run within the board's memory and timing constraints. This matters because the final firmware also includes a colour model, a movement model, sensor libraries, the Edge Impulse SDK, and serial JSON output. TinyML deployments need to fit within flash and RAM and must complete inference quickly enough that the user experience still feels live (Google, n.d.). The voice model was therefore evaluated not only by accuracy but also by whether it could run continuously on the board.",
    )
    add_figure(
        doc,
        SCREENSHOTS / "personal/voice_debug_bridge.png",
        "Fig. P7. Runtime bridge output proving the deployed voice model emits voice_debug and voice_start.",
        6.3,
    )
    p(
        doc,
        "After deployment, the firmware emitted voice_debug events containing the current state, threshold, arming status, streak, and classification scores. Figure P7 shows this runtime evidence. The implemented threshold was 0.75, with a startup arming delay to reduce immediate false triggers. For demo reliability the required streak was reduced to one confident detection, because earlier tests showed that demanding multiple consecutive high-confidence windows made the system miss spoken commands. This is a practical trade-off: the validation model looked capable, but a live room, microphone placement, and speech variation can still reduce confidence. The final design keeps the threshold high, adds an arming delay, and exposes debug scores so the operator can understand why the board did or did not advance.",
    )
    p(
        doc,
        "The main limitation of my voice component is that it is trained for a small vocabulary and a limited set of recording conditions. It should not be presented as general speech recognition. Its value is that it provides a focused, explainable trigger for a staged IoT workflow. Future improvements would include collecting more speakers, adding more negative samples from the actual demo room, testing different thresholds, and disabling debug output once the live behaviour is stable.",
    )


def architecture_section(doc: Document) -> None:
    heading(doc, "3. Group system architecture")
    p(
        doc,
        "The full group system is split into firmware, bridge, and web/cloud layers. Figure G1 shows the data flow and trust boundaries. The Arduino layer owns sensing and inference. It samples the microphone through the PDM library, reads colour data from the APDS9960 sensor, and reads acceleration and gyroscope data from the LSM9DS1 IMU. The firmware uses these inputs to run three TinyML models and prints JSON events to USB serial. The board intentionally does not connect to the network, which keeps the microcontroller firmware simpler and avoids putting server credentials on the device.",
    )
    add_figure(
        doc,
        ARCH_PATH,
        "Fig. G1. End-to-end architecture showing Arduino, serial bridge, web API, Postgres, Docker, and dashboard users.",
        6.5,
    )
    p(
        doc,
        "The bridge layer runs on the laptop connected to the board. It opens the serial port at 115200 baud, ignores non-JSON lines, validates that each line is a JSON object, and posts valid payloads to POST /api/movement. It also polls GET /api/bridge/control so the dashboard can request a clean bridge shutdown. This design was chosen because serial is reliable for the Nano board during a demo, while Python and HTTP are easier to debug than adding networking directly to the firmware.",
    )
    p(
        doc,
        "The web layer is a Next.js application. Route handlers implement the ingestion, latest-event, session-list, session-detail, complete-session, and bridge-control APIs. Next.js route handlers are designed for custom request handlers in the app directory (Next.js, n.d.). The dashboard polls the public read endpoints and renders the latest event, session state, timeline, and movement chart. When DATABASE_URL is configured, the event store persists data in PostgreSQL; without it, the same code falls back to memory for fast local testing. PostgreSQL was selected because session events and movement samples naturally fit relational tables with timestamps and foreign keys (PostgreSQL Global Development Group, n.d.).",
    )


def firmware_section(doc: Document) -> None:
    heading(doc, "4. Firmware and three-model integration")
    p(
        doc,
        "The firmware state machine has three main states: WAITING_FOR_VOICE, WAITING_FOR_COLOUR, and TRACKING_MOVEMENT. It starts in the voice state after reset. A successful start classification emits voice_start and moves to colour. A successful green classification emits colour_authenticated and moves to movement tracking. Once movement tracking starts, the board continuously sends movement events with raw IMU values and the movement model's predicted class and confidence.",
    )
    add_figure(
        doc,
        SCREENSHOTS / "arduino_state_machine_bridge_output_voice.png",
        "Fig. G3. Bridge output showing the voice phase and transition after start.",
        6.3,
    )
    p(
        doc,
        "Integrating three Edge Impulse exports was the hardest firmware issue. Each generated Arduino library normally expects to be the only Edge Impulse model in the sketch, so including three generated inferencing headers directly creates conflicting symbols such as run_classifier, run_classifier_init, model variables, and preprocessor macros. The solution was to create a local combined_inferencing Arduino library. It contains one SDK copy, merged model metadata, and explicit handles for the three models: voice model 970121, colour model 970107, and movement model 928825. The sketch includes only combined_inferencing.h and calls classifier functions through the selected handle.",
    )
    p(
        doc,
        "The colour model replaced the earlier scripted colour heuristic. The firmware reads APDS9960 red, green, blue, and clear values and maps them to the model input order ch1, ch2, ch3, and ch4. The model labels are blue, green, other, and red. The firmware requires the green class to exceed a 0.70 threshold for three consecutive reads before colour authentication succeeds. This is more robust than a single raw RGB threshold because the model can learn the teammate's recorded colour examples rather than relying only on a fixed ratio.",
    )
    add_figure(
        doc,
        SCREENSHOTS / "arduino_state_machine_bridge_output_colour.png",
        "Fig. G4. Bridge output showing colour authentication events.",
        6.3,
    )
    p(
        doc,
        "The movement model uses 88 accelerometer samples sampled at 44 Hz, with acceleration converted from g to m/s^2 before inference. The labels are down, idle, left, right, and up. The firmware maintains a rolling buffer, runs the model after the buffer is full, and emits movement JSON about every 250 ms. It includes both raw accelerometer and gyroscope values because raw telemetry is useful when debugging model mistakes, while movementClass and movementConfidence provide the TinyML interpretation.",
    )
    add_figure(
        doc,
        SCREENSHOTS / "arduino_state_machine_bridge_output_movement.png",
        "Fig. G5. Bridge output showing movement classification events.",
        6.3,
    )
    p(
        doc,
        "This firmware design keeps the demo explainable. Each state has one responsibility, and each state emits debug evidence while the demo is being stabilised. The trade-off is that the board state is one-way. Restarting the bridge does not reset the board, so a new full demo requires a board reset to return to WAITING_FOR_VOICE. This was accepted because it keeps the firmware logic simple and deterministic for a live demonstration.",
    )


def bridge_dashboard_section(doc: Document) -> None:
    heading(doc, "5. Bridge, dashboard, and data recording")
    p(
        doc,
        "The Python bridge is the boundary between embedded serial output and the web system. It uses pyserial to read the board, parses only lines that begin with JSON object syntax, and uses the requests library to POST each event. Verbose mode prints the HTTP status and compact JSON payload, which made it possible to debug the firmware sequence without opening the Arduino Serial Monitor. Only one process can own the serial port at a time, so the bridge and Serial Monitor are not run together.",
    )
    p(
        doc,
        "The dashboard is built around session recording rather than only the latest event. A setup_status event starts a new session. Later voice_start, colour_authenticated, and movement events attach to the active session. The database stores sessions, raw events, movement samples, and recording state. The movement_samples table separates chartable numeric telemetry from the raw JSON payload, while the events table preserves the original event for auditability. This gives the project both a live monitor and a persistent evidence store.",
    )
    add_figure(
        doc,
        SCREENSHOTS / "live_session_dashboard(18_full_demo_terminal).png",
        "Fig. G2. Live dashboard showing the current recorded session state.",
        4.2,
    )
    p(
        doc,
        "The dashboard shows business-facing status rather than only developer logs. It reports whether the voice, colour, and movement stages are waiting or complete, displays the raw latest JSON, and lists recorded sessions. The Stop current session button was added to control the lifecycle from the web UI. When the stop action is triggered, the server completes an authenticated run. If voice or colour authentication was not completed, the incomplete data is deleted because it does not represent a successful demo session. Stop also sets bridge-control state so the bridge exits cleanly instead of continuing to post movement samples.",
    )
    add_figure(
        doc,
        SCREENSHOTS / "dashboard_movement_chart.png",
        "Fig. G6. Recorded movement session visualised in the web dashboard.",
        6.4,
    )
    p(
        doc,
        "Figure G6 shows why recording sessions is more valuable than printing logs. The dashboard turns repeated movement events into a replayable chart with confidence values and timestamps. This supports post-demo analysis: the group can show that the system detected voice, authenticated colour, entered tracking, and captured movement data. For a coursework demo, this is important because it provides evidence even after the live interaction has ended.",
    )


def docker_section(doc: Document) -> None:
    heading(doc, "6. Cloud deployment and Docker")
    p(
        doc,
        "The deployable server path uses Docker Compose to run the web application and PostgreSQL database. Compose is designed for defining and running multi-container applications with a single configuration file (Docker, n.d.-a). In this project, the postgres service uses the postgres:16-alpine image and a named volume for persistent data, while the web service builds the Next.js application from web/Dockerfile. The web container runs the migration script before starting the production server, so a fresh server can create the required tables automatically.",
    )
    p(
        doc,
        "The server deployment intentionally excludes the Arduino and bridge from Docker. The bridge needs direct USB serial access to the board, so it stays on the laptop connected to the device and posts to the server URL. This separation is practical: the cloud/server part stores and visualises data, while the local laptop manages physical hardware. Docker Compose environment variables configure Postgres credentials, the public web port, and required API tokens (Docker, n.d.-b). The deployment script _update_server runs git pull --ff-only, rebuilds the containers, restarts the stack, and prints the service state.",
    )
    add_figure(
        doc,
        SCREENSHOTS / "docker_running.png",
        "Fig. G9. Docker Compose stack running the web application and database services.",
        5.8,
    )
    p(
        doc,
        "Postgres is not exposed broadly to the network. The Compose file binds the database port to 127.0.0.1 on the host, while the web container connects to it through the internal Docker network. This provides enough persistence and cloud-style deployment evidence for the project without turning the database into a public service. The default credentials are suitable only for demonstration and can be changed in .env before deployment.",
    )


def security_section(doc: Document) -> None:
    heading(doc, "7. Security using STRIDE")
    p(
        doc,
        "The project uses STRIDE as a lightweight security protocol and threat modelling method. STRIDE classifies threats as spoofing, tampering, repudiation, information disclosure, denial of service, and elevation of privilege (Microsoft, n.d.). For this system, the main assets are sensor events, session history, control actions, and shared secrets. The main trust boundaries are USB serial from the board to the bridge, HTTP from bridge/browser to the web API, and web API to PostgreSQL inside Docker.",
    )
    p(
        doc,
        "The implemented mitigation is role-based bearer-token protection. POST /api/movement and GET /api/bridge/control require BRIDGE_API_TOKEN when configured. POST /api/sessions/current/complete requires ADMIN_API_TOKEN when configured. Read-only dashboard endpoints remain public so a viewer can watch the demo without receiving write access. Bearer tokens are sent in the Authorization header, following the standard bearer-token pattern defined by RFC 6750 (IETF, 2012). On the server, token comparison uses Node.js crypto.timingSafeEqual to avoid simple timing leaks in equality checks (Node.js, n.d.).",
    )
    add_figure(
        doc,
        SCREENSHOTS / "security_unauthorized_401.png",
        "Fig. G7. Unauthorized API request returning HTTP 401 when no bearer token is supplied.",
        5.3,
    )
    add_figure(
        doc,
        SCREENSHOTS / "security_authorized_success.png",
        "Fig. G8. Authorized API request succeeding with a valid bearer token.",
        6.3,
    )
    p(
        doc,
        "Figures G7 and G8 show the difference between unauthorised and authorised requests. The goal is not enterprise identity management, but practical protection for a demo server. A random network user should not be able to fabricate board events, stop the bridge, or complete/delete sessions. The dashboard prompts for the admin token on the first protected stop action and stores it in browser localStorage for the demo operator.",
    )
    p(
        doc,
        "Residual risks remain. If the server is exposed directly over plain HTTP, tokens can be observed on the network; a real public deployment should use HTTPS behind a reverse proxy. Tokens are shared secrets rather than named user accounts, so anyone with the token has that role. Browser localStorage can be read by injected JavaScript, so the dashboard should not include untrusted scripts. Physical access to the board and bridge laptop is still trusted. These residual risks are acceptable for a university demonstration but should be addressed before a production deployment.",
    )


def testing_section(doc: Document) -> None:
    heading(doc, "8. Testing and evaluation")
    p(
        doc,
        "The manual test procedure starts the web application, starts the Python bridge, resets the board, and observes setup_status. The operator says start, shows green to the APDS9960 sensor, then moves the board. A successful run should record voice_start, colour_authenticated, and movement events in that order. The screenshots in this report provide evidence for each phase: voice transition in Figure G3, colour authentication in Figure G4, movement classification in Figure G5, and session replay in Figure G6.",
    )
    p(
        doc,
        "API tests were also used. Unauthorised mutation requests return 401 when tokens are configured, while authorised requests succeed. Docker testing checks that Compose starts both services and that migrations complete before the web server runs. The main evaluation result is that the system behaves as a full pipeline: local TinyML inference controls state transitions, the bridge reliably forwards JSON, the web app stores and visualises sessions, and security controls prevent unauthorised writes on the deployed path.",
    )


def conclusion(doc: Document) -> None:
    heading(doc, "9. Conclusion")
    p(
        doc,
        "The project demonstrates a complete TinyML IoT workflow rather than a disconnected model or a raw sensor dashboard. My voice recognition model provides the first interaction gate, and the group integration combines voice, colour, and movement models into a single firmware state machine. The Python bridge, Next.js dashboard, PostgreSQL storage, Docker deployment, and STRIDE security layer turn the board output into a deployable recorded-session system.",
    )
    p(
        doc,
        "The strongest business value is explainability. A stakeholder can watch the system progress from voice command to colour authentication to movement tracking and then review the captured session afterwards. Future work should improve the dataset size for all models, add HTTPS for public deployment, replace shared tokens with named users, and reduce debug output once the demo is stable. Even with those limitations, the current prototype shows how low-cost edge AI can be connected to a cloud-style evidence and control layer.",
    )


def references(doc: Document) -> None:
    heading(doc, "References")
    refs = [
        "Arduino (n.d.-a) Nano 33 BLE Sense. Available at: https://docs.arduino.cc/hardware/nano-33-ble-sense/ (Accessed: 28 April 2026).",
        "Arduino (n.d.-b) Arduino_APDS9960 library. Available at: https://github.com/arduino-libraries/Arduino_APDS9960 (Accessed: 28 April 2026).",
        "Arduino (n.d.-c) Arduino_LSM9DS1 library. Available at: https://github.com/arduino-libraries/Arduino_LSM9DS1 (Accessed: 28 April 2026).",
        "Docker (n.d.-a) Docker Compose documentation. Available at: https://docs.docker.com/compose/ (Accessed: 28 April 2026).",
        "Docker (n.d.-b) Interpolation: Docker Compose environment variables. Available at: https://docs.docker.com/compose/how-tos/environment-variables/variable-interpolation/ (Accessed: 28 April 2026).",
        "Docker (2022) How to use the Postgres Docker Official Image. Available at: https://www.docker.com/blog/how-to-use-the-postgres-docker-official-image/ (Accessed: 28 April 2026).",
        "Edge Impulse (n.d.-a) Run Arduino library. Available at: https://docs.edgeimpulse.com/docs/run-inference/arduino-library (Accessed: 28 April 2026).",
        "Edge Impulse (n.d.-b) Sound recognition. Available at: https://docs.edgeimpulse.com/docs/tutorials/audio-classification (Accessed: 28 April 2026).",
        "Google (n.d.) LiteRT for Microcontrollers. Available at: https://ai.google.dev/edge/litert/microcontrollers/overview (Accessed: 28 April 2026).",
        "IETF (2012) RFC 6750: The OAuth 2.0 Authorization Framework: Bearer Token Usage. Available at: https://www.rfc-editor.org/rfc/rfc6750 (Accessed: 28 April 2026).",
        "Mermaid (n.d.) Flowcharts syntax. Available at: https://mermaid.js.org/syntax/flowchart.html (Accessed: 28 April 2026).",
        "Microsoft (n.d.) Threats: Microsoft Threat Modeling Tool. Available at: https://learn.microsoft.com/en-us/azure/security/develop/threat-modeling-tool-threats (Accessed: 28 April 2026).",
        "Next.js (n.d.) Route Handlers. Available at: https://nextjs.org/docs/app/getting-started/route-handlers-and-middleware (Accessed: 28 April 2026).",
        "Node.js (n.d.) Crypto: crypto.timingSafeEqual. Available at: https://nodejs.org/api/crypto.html#cryptotimingsafeequala-b (Accessed: 28 April 2026).",
        "PostgreSQL Global Development Group (n.d.) PostgreSQL documentation. Available at: https://www.postgresql.org/docs/ (Accessed: 28 April 2026).",
    ]
    for ref in refs:
        para = doc.add_paragraph(ref)
        para.paragraph_format.first_line_indent = Cm(-0.55)
        para.paragraph_format.left_indent = Cm(0.55)
        para.paragraph_format.space_after = Pt(5)


def appendix(doc: Document) -> None:
    heading(doc, "Appendix A. Evidence map")
    add_table(
        doc,
        ["Evidence", "What it proves"],
        [
            ("Figures P1-P7", "Personal voice TinyML workflow, validation, on-device feasibility, and deployed runtime output."),
            ("Figures G1-G6", "End-to-end architecture, firmware state transitions, bridge forwarding, dashboard session recording, and movement replay."),
            ("Figures G7-G8", "STRIDE-inspired API protection, with unauthorised and authorised request outcomes."),
            ("Figure G9", "Docker Compose deployment path for the web app and PostgreSQL database."),
        ],
    )


def build_doc() -> None:
    generate_architecture_diagram()
    doc = Document()
    style_document(doc)
    add_page_number(doc.sections[0])
    cover(doc)
    executive_summary(doc)
    intro(doc)
    personal_section(doc)
    doc.add_section(WD_SECTION.NEW_PAGE)
    architecture_section(doc)
    doc.add_page_break()
    firmware_section(doc)
    bridge_dashboard_section(doc)
    docker_section(doc)
    security_section(doc)
    testing_section(doc)
    conclusion(doc)
    references(doc)
    appendix(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)
    print(ARCH_PATH)


if __name__ == "__main__":
    build_doc()
